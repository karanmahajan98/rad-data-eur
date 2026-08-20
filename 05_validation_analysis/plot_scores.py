"""
Generate figure for the R2 and RMSE scores of the shortwave and net radiation products.


Input : validation_tables_sw_netrad.docx
        Table 1 = net shortwave radiation (SSNR vs ERA5-Land SSR), FLUXNET
        Table 2 = net radiation (Rn vs ERA5-Land Rn), FLUXNET
Output: netrad_hybrid_figure.png 
"""
# import libraries
import numpy as np
import matplotlib.pyplot as plt
from docx import Document

#%%
DOCX_PATH = "/data/saved_vars/validation_tables_sw_netrad.docx"
OUTPUT_NAME = "/data/saved_vars/net_rad_hybrid_figure"

# Style features

# One colour per IGBP land-cover family. 
FAMILY_COLORS = {
    "Forest": "#217A3B",
    "Grassland": "#E8890C",
    "Cropland": "#7A29C9",
    "Shrub / savanna": "#6B3E1E",
    "Wetland / snow": "#4FB3E8",
}

# to map the raw IGBP land-cover code to one of the five families above.
IGBP_TO_FAMILY = {
    "ENF": "Forest", "DBF": "Forest", "EBF": "Forest", "MF": "Forest", "DNF": "Forest",
    "GRA": "Grassland",
    "CRO": "Cropland",
    "OSH": "Shrub / savanna", "CSH": "Shrub / savanna", "SAV": "Shrub / savanna", "WSA": "Shrub / savanna",
    "WET": "Wetland / snow", "SNO": "Wetland / snow",
}

COLOR_BETTER = "#1182D1"   # this study beats ERA5-Land
COLOR_WORSE = "#D92929"    # ERA5-Land beats this study

#%%

# read the two tables out of the Word document

def is_number(text):
    """Return True if `text` (a string) can be converted to a float."""
    try:
        float(text)
        return True
    except ValueError:
        return False
    
#%%

def read_site_table(table):
    """Read one FLUXNET validation table from the docx file.

    Each row of the table looks like:
        index, site_id, country, IGBP, ERA5_RMSE, ERA5_R2,
        ERA5_country_avg_RMSE, ERA5_country_avg_R2,
        product_RMSE, product_R2, product_country_avg_RMSE, product_country_avg_R2

    need columns 1-3 (site id, country, IGBP) and 4,5,8,9 (the two
    RMSE/R2 scores). The "country average" columns (6,7,10,11) are ignored
    here. country means are computed further down directly from the site
    rows, so they are always consistent with each other (see the note in
    `compute_country_means`).

    Parameters
    ----------
    table : docx.table.Table
        One table object from the Word document (doc.tables[0] or [1]).

    Returns
    -------
    list[dict]
        One dict per usable site, with keys:
        site_id (str), country (str), family (str),
        era5_rmse (float), era5_r2 (float), product_rmse (float), product_r2 (float)
    """
    sites = []
    current_country = None

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 12 or not cells[0].isdigit():
            continue  # this is a header row, skip it

        # The country column is merged across a block of rows in the Word
        # table, so it is only filled in on the first row of each country.
        # Carry the last seen value forward.
        if cells[2]:
            current_country = cells[2]

        era5_rmse, era5_r2, product_rmse, product_r2 = cells[4], cells[5], cells[8], cells[9]
        if not all(is_number(v) for v in (era5_rmse, era5_r2, product_rmse, product_r2)):
            continue  # site is missing a score (e.g. no ERA5-Land match) so  drop it

        igbp_code = cells[3]
        sites.append(dict(
            site_id=cells[1],
            country=current_country,
            family=IGBP_TO_FAMILY.get(igbp_code, "Wetland / snow"),
            era5_rmse=float(era5_rmse),
            era5_r2=float(era5_r2),
            product_rmse=float(product_rmse),
            product_r2=float(product_r2),
        ))
    return sites

#%% 
def compute_country_means(sites):
    """Average each score over the sites of each country.

    averages the same set of sites on both the ERA5-Land side and the product
    side. 

    Parameters
    ----------
    sites : list[dict]
        Output of `read_site_table`.

    Returns
    -------
    dict[str, dict]
        Maps country name to dict(era5_rmse, era5_r2, product_rmse,
        product_r2, n_sites), all floats except n_sites (int).
    """
    countries = sorted(set(site["country"] for site in sites))
    means = {}
    for country in countries:
        country_sites = [s for s in sites if s["country"] == country]
        means[country] = dict(
            era5_rmse=np.mean([s["era5_rmse"] for s in country_sites]),
            era5_r2=np.mean([s["era5_r2"] for s in country_sites]),
            product_rmse=np.mean([s["product_rmse"] for s in country_sites]),
            product_r2=np.mean([s["product_r2"] for s in country_sites]),
            n_sites=len(country_sites),
        )
    return means

#%%

# plot panels a & b product vs. ERA5-Land, one dot per site


def plot_score_comparison(ax, sw_sites, rn_sites, era5_key, product_key,
                           axis_limits, tick_values, tick_format,
                           panel_letter, panel_title):
    """Draw one 1:1 scatter panel: x = ERA5-Land score, y = this study's score.


    Parameters
    ----------
    ax : matplotlib.axes.Axes
        The subplot to draw into.
    sw_sites, rn_sites : list[dict]
        Site records from `read_site_table`, for the shortwave (SW) and net
        radiation (Rn) tables respectively.
    era5_key, product_key : str
        Which score to plot, e.g. "era5_rmse" / "product_rmse" or
        "era5_r2" / "product_r2".
    axis_limits : tuple[float, float]
        (min, max) shared by both the x- and y-axis (the plot is square).
    tick_values : list[float]
        Where to put axis ticks.
    tick_format : str
        printf-style format for the tick labels, e.g. "%d" or "%.1f".
    panel_letter : str
        Panel label shown above the plot, e.g. "a".
    panel_title : str
        Axis/panel title shown next to the letter, e.g. "RMSE (W m$^{-2}$)".

    Returns
    -------
    None
    """
    axis_min, axis_max = axis_limits
    ax.set_xlim(axis_min, axis_max)
    ax.set_ylim(axis_min, axis_max)

    # 1:1 reference line
    ax.plot(axis_limits, axis_limits, color="gray", linewidth=0.8, zorder=1)

    # SSNR = filled circles, Rn = hollow diamonds. Both coloured by IGBP.
    for site in sw_sites:
        color = FAMILY_COLORS[site["family"]]
        ax.plot(site[era5_key], site[product_key], marker="o", markersize=5,
                markerfacecolor=color, markeredgecolor="white", markeredgewidth=0.5, zorder=3)
    for site in rn_sites:
        color = FAMILY_COLORS[site["family"]]
        ax.plot(site[era5_key], site[product_key], marker="D", markersize=5,
                markerfacecolor="none", markeredgecolor=color, markeredgewidth=1.1, zorder=3)

    ax.set_xticks(tick_values)
    ax.set_yticks(tick_values)
    ax.set_xticklabels([tick_format % t for t in tick_values])
    ax.set_yticklabels([tick_format % t for t in tick_values])
    ax.set_xlabel("ERA5-Land")
    ax.set_ylabel("This study")
    ax.set_title(f"{panel_letter}  {panel_title}", loc="left", fontweight='bold')

    # which side is better labels in the corners of the plot.
    if "rmse" in era5_key:
        ax.text(0.03, 0.97, "worse than ERA5-Land", color=COLOR_WORSE, fontsize=10,
                ha="left", va="top", transform=ax.transAxes)
        ax.text(0.97, 0.03, "better than ERA5-Land", color=COLOR_BETTER, fontsize=10,
                ha="right", va="bottom", transform=ax.transAxes)
    else:
        ax.text(0.03, 0.97, "better than ERA5-Land", color=COLOR_BETTER, fontsize=10,
                ha="left", va="top", transform=ax.transAxes)
        ax.text(0.97, 0.03, "worse than ERA5-Land", color=COLOR_WORSE, fontsize=10,
                ha="right", va="bottom", transform=ax.transAxes)

#%%

# plot panel c -- country-mean change relative to ERA5-Land


def plot_country_deltas(ax, countries, country_means, sites,
                         era5_key, product_key, half_range, decimals,
                         column_title, show_country_labels):
    """Draw one column of panel c: a horizontal bar per country showing the
    mean change (this study minus ERA5-Land) in one metric, with the
    individual FLUXNET towers plotted as small dots behind the bars.".

    Parameters
    ----------
    ax : matplotlib.axes.Axes
        The subplot to draw into (one of the four bottom-row columns).
    countries : list[str]
        All country names to show, top-to-bottom, shared across all four
        columns so the rows line up.
    country_means : dict[str, dict]
        Output of `compute_country_means` for this variable (SW or Rn).
    sites : list[dict]
        Site records for this variable, used to plot the individual-tower
        dots.
    era5_key, product_key : str
        Which pair of scores to compare, e.g. "era5_rmse" / "product_rmse".
    half_range : float
        The bar axis runs from -half_range to +half_range; values further
        out are clipped to the edge (the true value is still in the label).
    decimals : int
        How many decimal places to print in the value labels.
    column_title : str
        Title shown above this column, e.g. r"$\\Delta$RMSE, SSNR".
    show_country_labels : bool
        Only the left-most column needs the country names as y-axis labels;
  
    Returns
    -------
    None
    """
    rng = np.random.default_rng(seed=0)  # so re-runs look identical, used to shift the dots in y direction so they don't all overlap
    y_positions = np.arange(len(countries))

    for y, country in zip(y_positions, countries):
        stats = country_means.get(country)
        if stats is None:
            ax.text(half_range * 1.12, y, "n/a", ha="left", va="center",
                     fontsize=9, color="gray", clip_on=False)
            continue

        delta = stats[product_key] - stats[era5_key]
        this_study_is_better = delta < 0 if "rmse" in era5_key else delta > 0
        bar_color = COLOR_BETTER if this_study_is_better else COLOR_WORSE

        clipped_delta = np.clip(delta, -half_range, half_range)
        ax.barh(y, clipped_delta, height=0.6, color=bar_color, alpha=0.9, zorder=2)

        # Individual tower dots, shifter slightly in y so they don't all overlap on one line.
        country_sites = [s for s in sites if s["country"] == country]
        tower_deltas = np.array([s[product_key] - s[era5_key] for s in country_sites])
        tower_deltas = np.clip(tower_deltas, -half_range, half_range)
        jitter = rng.uniform(-0.25, 0.25, size=len(tower_deltas))
        ax.scatter(tower_deltas, y + jitter, s=6, color="black", alpha=0.35, zorder=3)

#        sign = "+" if delta > 0 else "−"
#        ax.text(half_range * 1.12, y, f"{sign}{abs(delta):.{decimals}f}",
#                ha="left", va="center", fontsize=8, clip_on=False)

        ax.text(half_range * 1.12, y,
                f"({stats[product_key]:.2f}, {stats[era5_key]:.2f})",
                ha="left", va="center", fontsize=9, clip_on=False)

    ax.axvline(0, color="#c8c3bb", linewidth=0.8, zorder=1)
    # a little extra room on the left and the right for the labels
    ax.set_xlim(-half_range*1.1, half_range * 1.45)
    ax.set_xticks([-half_range, 0, half_range])
    ax.set_ylim(-0.5, len(countries) - 0.5)
    ax.invert_yaxis()  # first country at the top, matching a table layout
    ax.set_title(column_title, fontsize=11)

    if show_country_labels:
        labels = [f"{c}" for c in countries]
        ax.set_yticks(y_positions)
        ax.set_yticklabels(labels, fontsize=10)
    else:
        ax.set_yticks(y_positions)
        ax.set_yticklabels([])

    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.tick_params(left=False)

#%%

# build the figure

def make_legend_handles():
    """Build the matplotlib legend handles shared by the top-row panels:
    one for each product marker and one per land-cover class.

    Returns
    -------
    list[matplotlib.lines.Line2D]
    """
    handles = [
        plt.Line2D([0], [0], marker="o", linestyle="none", markersize=9,
                   markerfacecolor="#54524e", markeredgecolor="none", label="SSNR"),
        plt.Line2D([0], [0], marker="D", linestyle="none", markersize=9,
                   markerfacecolor="none", markeredgecolor="#54524e", markeredgewidth=1.2, label="$R_n$"),
    ]
    for family, color in FAMILY_COLORS.items():
        handles.append(plt.Line2D([0], [0], marker="o", linestyle="none", markersize=9,
                                   markerfacecolor=color, markeredgecolor="none", label=family))
    return handles

#%%
def main():
    document = Document(DOCX_PATH)
    sw_sites = read_site_table(document.tables[0])
    rn_sites = read_site_table(document.tables[1])
    sw_country_means = compute_country_means(sw_sites)
    rn_country_means = compute_country_means(rn_sites)

    print(f"SW sites: {len(sw_sites)}   Rn sites: {len(rn_sites)}")
    print(f"SW countries: {len(sw_country_means)}   Rn countries: {len(rn_country_means)}")

    all_countries = sorted(set(sw_country_means) | set(rn_country_means))

    # Figure layout: a 2-row grid. Row 1 holds the two scatter panels
    # (a, b); row 2 holds panel c split into four columns, one per metric.
    # The two rows are built from separate GridSpecs (rather than one grid
    # shared across both) so each can use its own column spacing 
    
    #There is a large gap between the two rows (top row ends at 0.66, bottom row
    # starts at 0.48) so the top row's legend can sit in between, below
    # the scatter panels instead of above them.
    fig = plt.figure(figsize=(11, 13))
    ROW_LEFT, ROW_RIGHT = 0.06, 0.98
    grid_top = fig.add_gridspec(nrows=1, ncols=2, left=ROW_LEFT, right=ROW_RIGHT,
                                 top=0.93, bottom=0.66, wspace=0.28)
    GRID_BOTTOM_TOP, GRID_BOTTOM_BOTTOM = 0.48, 0.08
    grid_bottom = fig.add_gridspec(nrows=1, ncols=4, left=ROW_LEFT, right=ROW_RIGHT,
                                    top=GRID_BOTTOM_TOP, bottom=GRID_BOTTOM_BOTTOM,
                                    wspace=0.45)

    ax_rmse = fig.add_subplot(grid_top[0, 0])
    ax_r2 = fig.add_subplot(grid_top[0, 1])

    plot_score_comparison(
        ax_rmse, sw_sites, rn_sites, "era5_rmse", "product_rmse",
        axis_limits=(10, 85), tick_values=[10, 20, 30, 40, 50, 60, 70, 80],
        tick_format="%d", panel_letter="a)", panel_title="RMSE (W m$^{\mathbf{-2}}$)")
    plot_score_comparison(
        ax_r2, sw_sites, rn_sites, "era5_r2", "product_r2",
        axis_limits=(0.1, 1.0), tick_values=[0.2, 0.4, 0.6, 0.8, 1.0],
        tick_format="%.1f", panel_letter="b)", panel_title="Coefficient of determination, $\mathbf{R^2}$")

    # Legend for panels a/b, placed in the gap below them

    fig.legend(handles=make_legend_handles(), loc="center",
               bbox_to_anchor=(ROW_LEFT, 0.55, ROW_RIGHT - ROW_LEFT, 0.04), 
               ncol=7, mode="expand", frameon=False, fontsize=10)

    country_axes = [fig.add_subplot(grid_bottom[0, i]) for i in range(4)]
    column_specs = [
        # (era5_key, product_key, half_range, decimals, title, means, sites)
        ("era5_rmse", "product_rmse", 10.0, 1, r"$\Delta$RMSE, SSNR", sw_country_means, sw_sites),
        ("era5_r2", "product_r2", 0.10, 3, r"$\Delta R^2$, SSNR", sw_country_means, sw_sites),
        ("era5_rmse", "product_rmse", 14.0, 1, r"$\Delta$RMSE, $R_n$", rn_country_means, rn_sites),
        ("era5_r2", "product_r2", 0.25, 3, r"$\Delta R^2$, $R_n$", rn_country_means, rn_sites),
    ]
    for i, (ax, (era5_key, product_key, half_range, decimals, title, means, sites)) in enumerate(
            zip(country_axes, column_specs)):
        plot_country_deltas(ax, all_countries, means, sites, era5_key, product_key,
                             half_range, decimals, title, show_country_labels=(i == 0))

    fig.text(0.05, GRID_BOTTOM_TOP + 0.045, "c)  Country-mean change relative to ERA5-Land",
             fontsize=12, fontweight="bold")

    legend_handles = [
        plt.Rectangle((0, 0), 1, 1, color=COLOR_BETTER, label="this study better"),
        plt.Rectangle((0, 0), 1, 1, color=COLOR_WORSE, label="ERA5-Land better"),
        plt.Line2D([0], [0], marker="o", linestyle="none", color="black", alpha=0.35,
                   markersize=5, label="individual towers"),
    ]
    fig.legend(handles=legend_handles, loc="lower center",
               bbox_to_anchor=(0.5, 0.0), ncol=3, frameon=False, fontsize=10)

    fig.savefig(f"{OUTPUT_NAME}.png", dpi=300, facecolor="white", bbox_inches="tight")

    print(f"wrote {OUTPUT_NAME}.png")


if __name__ == "__main__":
    main()